import os
import textract
from tkinter import filedialog, messagebox
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from reportlab.lib.styles import ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from win32com.client import Dispatch

# This code is mainly for reading and outputing files only


# Read file
def importFile(path):
    if not os.path.exists(path):
        return ''

    txt = ''
    if path.lower().endswith(('.docx', '.pdf')):
        rep = (b'\n\n', b'\n') if path.lower().endswith(
            ('.doc', '.docx')) else (b'\r\n', b'\n')
        try:
            txt = textract.process(path).replace(*rep).decode('utf-8')
        except textract.exceptions.MissingFileError:
            messagebox.showerror(
                'Error', 'The file cannot be found, please check your file location.')
            return ''
        except:
            messagebox.showerror('Error', 'Unknown error.')
            return ''
    elif path.lower().endswith('.doc'):
        # Open word application
        word = Dispatch('Word.Application')

        # Make it work in background
        word.Visible = False

        # Prevent the saving alerts
        word.DisplayAlerts = False
        try:
            doc = word.Documents.Open(FileName=path, Encoding='utf-8')
            for p in doc.paragraphs:
                txt += p.Range.Text
                print(p.Range.Text)
        except:
            messagebox.showerror('Error', 'Unknown error.')
            word.Quit
            return ''
        txt = txt.replace('\r', '')
        print('.doc txt:', txt)
        doc.Close()
        word.Quit
    else:
        try:
            f = open(path, 'r', encoding='utf-8')
            txt = f.read()
        except:
            messagebox.showerror('Error', 'Unknown error.')
            return ''

    return txt


# Output to a file
def exportFile(output, underlineOnOff, strikethroughOnOff, highlightOnOff):
    if not output:
        messagebox.showwarning(
            'ReFile', 'There is no text in the output, please add text before saving output.')
        return

    print('output:', output)

    try:
        # May add HTML format
        file = filedialog.asksaveasfile(initialdir='/', initialfile='output.docx', defaultextension='*.*', filetypes=[
            ('All Files', '.doc .docx .pdf .txt'),
            ('Word Documents', '.doc .docx'),
            ('Adobe PDF', '.pdf')])
    except PermissionError:
        messagebox.showwarning(
            'Error', 'Cannot save the output to a file, please close the file or see the permisson of the file if you are overwriting it.')
        return
    except:
        messagebox.showerror('Error', 'Unknown error.')
        return

    if file is None:
        return

    if file.name.lower().endswith(('.doc', '.docx')):
        # Create a template
        document = Document()

        # Add a paragraph
        p = document.add_paragraph()

        for action, text in output:
            # Add text to paragraph reference
            run = p.add_run(text)
            run.font.size = Pt(14) if action == 0 else Pt(16)

            # Get the XML tag
            tag = run._r

            # Create an XML element
            shd = OxmlElement('w:shd')

            # Add attributes to the element
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')

            run.font.strike = strikethroughOnOff and action == -1
            run.font.underline = underlineOnOff and action == 1

            # Set the shading colour
            shd.set(qn('w:fill'), 'FFAAAA' if action == -1 else 'AAFFAA')

            # Append the element under <w:rPr>
            if action != 0:
                tag.rPr.append(shd)

            document.save(file.name)
    elif file.name.lower().endswith('.pdf'):
        # Change the text for outputing a PDF
        pdf_text = ''

        # Set both underline and strike tag
        underline_open_tag, underline_close_tag = (
            '<u>', '</u>') if underlineOnOff else ('', '')
        strike_open_tag, strike_close_tag = (
            '<strike>', '</strike>') if strikethroughOnOff else ('', '')

        # Replace newline character with br tag
        rep = ('\n', '<br/>')
        for action, text in output:
            deletion_colour = ' backcolor="#FFAAAA" ' if action == -1 and highlightOnOff else ' '
            insertion_colour = ' backcolor="#AAFFAA" ' if action == 1 and highlightOnOff else ' '
            if action == -1:
                pdf_text += '<span' + deletion_colour + 'fontSize=16>' + \
                    strike_open_tag + \
                    text.replace(*rep) + strike_close_tag + '</span>'
            elif action == 1:
                pdf_text += '<span' + insertion_colour + 'fontSize=16>' + \
                    underline_open_tag + \
                    text.replace(*rep) + underline_close_tag + '</span>'
            else:
                pdf_text += '<span>' + text.replace(*rep) + '</span>'

        # Register a font
        pdfmetrics.registerFont(
            TTFont('Noto Sans HK', './font/NotoSansHK-Regular.ttf'))

        # Create a style
        style = ParagraphStyle(
            name='Normal',
            fontName='Noto Sans HK',
            textColor='black',
            fontSize=14,
            leading=21,
        )

        # Output PDF file
        simpledoctemplate = SimpleDocTemplate(file.name)
        simpledoctemplate.build([Paragraph(pdf_text, style)])
    elif file.name.lower().endswith('.txt'):
        with open(file.name, 'w', encoding='utf-8') as f:
            f.write(''.join([word for _, word in output]))
    else:
        messagebox.showwarning(
            'Warning', 'The file format you are saving may not be compatible with text files, please try again.')

    print('Output path:', file.name)
    file.close()
